"""
Word/DOCX 解析器

识别 Heading 1-9 标题样式，自动映射层级；
识别加粗、高亮文本并标记为重点。
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import ParsedDocument


def parse_docx(file_path: str, priority: str = "normal") -> ParsedDocument:
    """解析 DOCX 文件，返回 ParsedDocument"""
    doc = ParsedDocument(
        file_path=file_path,
        file_type="docx",
        priority=priority,
    )

    try:
        document = Document(file_path)
    except Exception as e:
        doc.parse_errors.append(f"无法打开 DOCX 文件: {e}")
        return doc

    doc.total_pages = 1  # DOCX 没有页数概念

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 识别标题层级
        level = _get_heading_level(para)

        # 回退：无样式时用文本模式推断层级
        if level == 0:
            level = _infer_level_from_text(text)

        # 识别重点文本（加粗/高亮/下划线）
        is_important = _has_format_marking(para)

        # 检测页眉页脚类内容（通常被放主文档流，但可通过格式排除）
        if _is_header_footer_like(para, text):
            continue

        doc.add_paragraph(
            text,
            level=level,
            is_important=is_important,
            raw_style=para.style.name if para.style else "Normal",
        )

    # 也处理表格中的文本
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and len(cell_text) > 3:
                    doc.add_paragraph(
                        cell_text,
                        level=0,
                        raw_style="TableCell",
                    )

    return doc


def _get_heading_level(para) -> int:
    """从段落样式中提取标题层级"""
    if para.style is None:
        return 0

    style_name = para.style.name.lower()

    # 标准 Heading 1-9
    for i in range(1, 10):
        if f"heading {i}" in style_name or f"heading{i}" in style_name:
            return i

    # 中文标题样式
    heading_map = {
        "标题 1": 1, "标题 2": 2, "标题 3": 3,
        "标题 4": 4, "标题 5": 5,
        "标题1": 1, "标题2": 2, "标题3": 3,
        "title": 1, "subtitle": 2,
    }
    for key, level in heading_map.items():
        if key in style_name:
            return level

    # Outline level (Word 段落属性)
    try:
        pf = para.paragraph_format
        if pf.outline_level is not None and pf.outline_level < 9:
            return pf.outline_level + 1
    except Exception:
        pass

    return 0


def _has_format_marking(para) -> bool:
    """检测段落中是否有加粗、高亮、下划线标记"""
    for run in para.runs:
        if run.bold:
            return True
        if run.underline:
            return True
        if run.font.highlight_color is not None:
            return True
    return False


def _is_header_footer_like(para, text: str) -> bool:
    """检测页眉页脚类内容"""
    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT and len(text) < 6:
        return True  # 右侧短文本 → 可能是页码
    if text.isdigit() and len(text) <= 4:
        return True   # 纯数字 4 位以内 → 可能是页码
    if "页眉" in para.style.name.lower() or "页脚" in para.style.name.lower():
        return True
    if "header" in para.style.name.lower() or "footer" in para.style.name.lower():
        return True
    return False


def _infer_level_from_text(text: str) -> int:
    """回退：从纯文本模式推断标题层级（无样式 DOCX 用）"""
    import re

    # "第X章" / "第X节" → level 1
    if re.match(r"^第[一二三四五六七八九十\d]+[章节]", text):
        return 1

    # "一、" / "二、" → level 1
    if re.match(r"^[一二三四五六七八九十]+[、．]", text):
        return 1

    # "1." / "1)" / "1、" → level 2
    if re.match(r"^\d+[\.\、\)]", text):
        return 2

    # "1.1" / "1.1.1" → level 取决于点数
    m = re.match(r"^(\d+\.)+", text)
    if m:
        dots = m.group().count(".")
        return min(dots + 1, 5)

    # "（一）" → level 2
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return 2

    # 短行（<20字）不以句号/逗号结尾 → 可能是标题
    if len(text) < 20 and not text.endswith(("。", "，", "；", ".", ",", ";")):
        return 1

    return 0
